import os

os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'

import json

import subprocess
import sys

import tensorflow
import warnings
import pickle
import numpy as np
import librosa
import matplotlib.pyplot as plt

from vosk import Model, KaldiRecognizer, SetLogLevel
from pydub import AudioSegment

import docx
from docx.shared import RGBColor


tensorflow.keras.config.disable_interactive_logging()
warnings.simplefilter("ignore")
warnings.filterwarnings("ignore", category=DeprecationWarning)

class SplitWavAudioMubin():
    def __init__(self, folder, filename):
        self.folder = folder
        self.filename = filename
        self.filepath = folder + '/' + filename

        self.audio = AudioSegment.from_wav(self.filepath)

    def get_duration(self):
        return self.audio.duration_seconds

    def single_split(self, from_msec, to_msec, split_filename):
        t1 = from_msec
        t2 = to_msec
        split_audio = self.audio[t1:t2]
        split_audio.export(self.folder + '/' + split_filename, format="wav")


def extract_features(data, sample_rate):
    #     # ZCR
    result = np.array([])
    # zcr = np.mean(librosa.feature.zero_crossing_rate(y=data).T, axis=0)
    # result = np.hstack((result, zcr))  # stacking horizontally

    # Chroma_stft
    stft = np.abs(librosa.stft(data))
    chroma_stft = np.mean(librosa.feature.chroma_stft(S=stft, sr=sample_rate).T, axis=0)
    result = np.hstack((result, chroma_stft))  # stacking horizontally
    #
    # MFCC
    # print(path)
    mfcc = np.mean(librosa.feature.mfcc(y=data, sr=sample_rate).T, axis=0)
    result = np.hstack((result, mfcc))  # stacking horizontally

    # Root Mean Square Value
    rms = np.mean(librosa.feature.rms(y=data).T, axis=0)
    result = np.hstack((result, rms))  # stacking horizontally

    # MelSpectogram
    mel = np.mean(librosa.feature.melspectrogram(y=data, sr=sample_rate).T, axis=0)
    result = np.hstack((result, mel))  # stacking horizontally
    #
    return result


encoder = pickle.load(open("encoder.pkl", "rb"))
scaler = pickle.load(open("scaler.pkl", "rb"))

SAMPLE_RATE = 16000
FILE_PATH = "audio/"
FILE_NAME = sys.argv[1]
SetLogLevel(0)

statistic = {0:0,
             1:0,
            2:0,
            3:0,
            4:0,
            5:0}

colors = {  0: RGBColor(171, 37, 37),
            1:RGBColor(111, 79, 40),
            2:RGBColor(255,164,33),
            3:RGBColor(72,164,63),
            4:RGBColor(157,163,166),
            5:RGBColor(52,129,184)}

use_colors = sys.argv[2] == "true"
gen_statistic = sys.argv[3] == "true"

model = Model(lang="ru", model_name="vosk-model-ru-0.42")

rec = KaldiRecognizer(model, SAMPLE_RATE)
rec.SetWords(True)
rec.SetPartialWords(True)

kmodel = tensorflow.keras.models.load_model("model.keras")


def make_predict(file_name, start, end):
    folder = "audio"
    file = f"cuted{start}-{end}.wav"
    cut = SplitWavAudioMubin(folder, file_name)
    cut.single_split(start, end, file)

    subprocess.Popen(
        f"ffmpeg -i audio/{file} -ar 16000 -ac 1 -acodec pcm_s16le audio/F{file}".split(),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE).wait()

    data, sample_rate = librosa.load(f"audio/F{file}")
    data = extract_features(data, sample_rate)
    data = scaler.transform(np.array([data]))
    data = np.expand_dims(data, axis=2)

    pr = kmodel.predict(data)

    os.system(f"rm -rf audio/{file}")
    os.system(f"rm -rf audio/F{file}")

    return np.argmax(pr)

def add_to_word(text, color):
    font = paragraph.add_run(text).font
    if color != -1:
        font.color.rgb = colors[color]
    paragraph.add_run(" ")


doc = docx.Document()
paragraph = doc.add_paragraph("")

with subprocess.Popen(["ffmpeg", "-loglevel", "quiet", "-i",
                       FILE_PATH + FILE_NAME,
                       "-ar", str(SAMPLE_RATE), "-ac", "1", "-f", "s16le", "-"],
                      stdout=subprocess.PIPE) as process:
    while True:
        data = process.stdout.read(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            data1 = json.loads(rec.FinalResult())

            i = 0
            if 'result' in data1:
                word_count = len(data1["result"])


                while i + 4 < word_count:
                    start = int(data1["result"][i]["start"] * 1000)
                    end = int(data1["result"][i + 4]["end"] * 1000)

                    sent = " ".join([data1["result"][j]["word"] for j in range(i, i+5)])

                    if i == 0:
                        sent = sent.capitalize()
                    if i+5 == word_count:
                        sent += "."

                    pr = make_predict(FILE_NAME, start, end)
                    if use_colors:
                        add_to_word(sent, pr)
                        print(sent, "@", pr, sep="")
                    else:
                        add_to_word(sent, -1)
                        print(sent, "@", sep="")

                    i += 5
                    statistic[pr] += 5

                if i < word_count:
                    start = int(data1["result"][i]["start"] * 1000)
                    end = int(data1["result"][word_count-1]["end"] * 1000)
                    sent = " ".join([data1["result"][j]["word"] for j in range(i, word_count)]) + "."

                    if i == 0:
                        sent = sent.capitalize()

                    pr = make_predict(FILE_NAME, start, end)

                    if use_colors:
                        add_to_word(sent, pr)
                        print(sent, "@", pr, sep="")
                    else:
                        add_to_word(sent, -1)
                        print(sent, "@", sep="")

                    statistic[pr] += 1

    data1 = json.loads(rec.FinalResult())
    if 'result' in data1:
        start = int(data1["result"][0]["start"] * 1000)
        end = int(data1["result"][-1]["end"] * 1000)

        sent = " ".join([data1["result"][j]["word"] for j in range(0, len(data1["result"]))])
        pr = make_predict(FILE_NAME, start, end)

        if use_colors:
            add_to_word(sent, pr)
            print(sent, "@", pr, sep="")
        else:
            add_to_word(sent, -1)
            print(sent, "@", sep="")

        statistic[pr] += len(data1["result"])

    names = ['angry', 'disgust', 'inspiration', 'fun', 'neutral', 'sad']
    values = [statistic[i] for i in statistic]

    if gen_statistic:
        plt.figure(figsize=(9, 3))
        plt.bar(names, values)
        plt.savefig(fname=f'statistic/{FILE_NAME[:-4]}.png')

    doc.save(f'./files/{FILE_NAME[:-4]}.docx')
    # print(f"{statistic[0]} {statistic[1]} {statistic[2]} {statistic[3]} {statistic[4]} {statistic[5]}@")